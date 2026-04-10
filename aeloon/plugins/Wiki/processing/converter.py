"""Local document converters used by the wiki digest flow."""

from __future__ import annotations

import csv
from pathlib import Path
from typing import Protocol, runtime_checkable


@runtime_checkable
class DocumentConverter(Protocol):
    """Protocol for source-to-markdown converters."""

    async def convert(self, source_path: Path) -> str: ...


class PdfConverter:
    """PDF to Markdown via PyMuPDF or pdfminer.six."""

    async def convert(self, source_path: Path) -> str:
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(str(source_path))
            parts: list[str] = []
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    parts.append(text)
            doc.close()
            markdown = "\n\n".join(parts)
            if markdown.strip():
                return markdown
        except ImportError:
            pass

        try:
            from pdfminer.high_level import extract_text

            text = extract_text(str(source_path))
            if text.strip():
                return text
        except ImportError:
            pass

        raise RuntimeError(
            "PDF conversion requires pdfminer.six or PyMuPDF. "
            "Install with: pip install pdfminer.six or pip install PyMuPDF"
        )


class DocxConverter:
    """DOCX to Markdown via python-docx."""

    async def convert(self, source_path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError(
                "DOCX conversion requires python-docx. pip install python-docx"
            ) from exc

        doc = DocxDocument(str(source_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if style.startswith("Heading"):
                try:
                    level = int(style.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                parts.append(f"{'#' * min(level, 6)} {text}")
            else:
                parts.append(text)

        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)
            if rows:
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
                body = "\n".join("| " + " | ".join(row) + " |" for row in rows[1:])
                parts.extend([header, separator, body])

        return "\n\n".join(parts)


class MarkdownConverter:
    """Pass through Markdown content with tolerant decoding."""

    async def convert(self, source_path: Path) -> str:
        return source_path.read_text(encoding="utf-8", errors="replace")


class TextConverter:
    """Plain text to Markdown."""

    async def convert(self, source_path: Path) -> str:
        return source_path.read_text(encoding="utf-8", errors="replace")


class CsvConverter:
    """CSV to a Markdown table."""

    async def convert(self, source_path: Path) -> str:
        rows: list[list[str]] = []
        with source_path.open(encoding="utf-8", errors="replace", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                rows.append(row)

        if not rows:
            return ""

        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
        body = "\n".join("| " + " | ".join(row) + " |" for row in rows[1:])
        return f"{header}\n{separator}\n{body}"


class ConverterFactory:
    """Dispatch document conversion by file extension."""

    def __init__(self) -> None:
        self._converters: dict[str, DocumentConverter] = {
            "pdf": PdfConverter(),
            "docx": DocxConverter(),
            "md": MarkdownConverter(),
            "txt": TextConverter(),
            "csv": CsvConverter(),
        }

    def get(self, ext: str) -> DocumentConverter:
        key = ext.lower()
        if key not in self._converters:
            raise ValueError(f"Unsupported format: .{key}")
        return self._converters[key]

    async def convert(self, source_path: Path) -> str:
        converter = self.get(source_path.suffix.lstrip("."))
        return await converter.convert(source_path)
